#!/usr/bin/env python3
"""Generate eval cases for the public enterprise project-document corpus."""

from __future__ import annotations

import argparse
import json
import os
import random
import re
import sys
import time
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import httpx
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from dotenv import dotenv_values
from pypdf import PdfReader


REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_CORPUS_DIR = REPO_ROOT / "benchmarks" / "enterprise_project_docs"
DEFAULT_RAW_ROOT = DEFAULT_CORPUS_DIR
DEFAULT_OUTPUT = DEFAULT_CORPUS_DIR / "eval_cases_sample_50.json"
DEFAULT_MODEL = "gemini-2.5-pro"
VERSION = "enterprise_project_docs_200_sample_v1"
MAX_DOC_CHARS = 8000

SAMPLE_QUOTAS = {
    "fact_qa": 16,
    "entity_relation": 12,
    "multi_hop": 12,
    "summary": 8,
    "distractor": 2,
}

CASE_PREFIX = {
    "fact_qa": "fact",
    "entity_relation": "relation",
    "multi_hop": "multi",
    "summary": "summary",
    "distractor": "distractor",
}


@dataclass
class GenerationStats:
    prompt_tokens: int = 0
    output_tokens: int = 0
    total_tokens: int = 0


def load_env_key(env_path: Path) -> str:
    values = dotenv_values(env_path) if env_path.exists() else {}
    key = os.environ.get("GEMINI_API_KEY") or values.get("GEMINI_API_KEY") or ""
    if key.strip():
        return key.strip()

    placeholder = "GEMINI_API_KEY=填入你的 Gemini API Key\n"
    if env_path.exists():
        current = env_path.read_text(encoding="utf-8")
        if "GEMINI_API_KEY=" not in current:
            with env_path.open("a", encoding="utf-8") as handle:
                if current and not current.endswith("\n"):
                    handle.write("\n")
                handle.write(placeholder)
    else:
        env_path.write_text(placeholder, encoding="utf-8")
    raise RuntimeError("GEMINI_API_KEY is missing. Added a placeholder to .env; fill it before running.")


def load_manifest(path: Path) -> list[dict[str, Any]]:
    manifest = json.loads(path.read_text(encoding="utf-8"))
    documents = manifest if isinstance(manifest, list) else manifest.get("documents", [])
    if len(documents) != 200:
        raise ValueError(f"Expected 200 documents, got {len(documents)} from {path}")
    return documents


def read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages)


def read_table(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        frame = pd.read_csv(path, nrows=80)
        return frame.to_csv(index=False)

    sheets = pd.read_excel(path, sheet_name=None, nrows=60)
    chunks = []
    for sheet_name, frame in list(sheets.items())[:4]:
        chunks.append(f"[sheet: {sheet_name}]")
        chunks.append(frame.to_csv(index=False))
    return "\n".join(chunks)


def read_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text("\n", strip=True)


def read_docx(path: Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = []
    for table in document.tables[:10]:
        for row in table.rows[:30]:
            tables.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + tables)


def read_doc(path: Path) -> str:
    try:
        return read_docx(path)
    except Exception:
        return path.read_bytes().decode("utf-8", errors="ignore")


def read_document(raw_root: Path, doc: dict[str, Any]) -> str:
    path = raw_root / doc["local_path"]
    suffix = path.suffix.lower()
    if not path.exists():
        raise FileNotFoundError(path)
    if suffix == ".pdf":
        text = read_pdf(path)
    elif suffix in {".xlsx", ".xls", ".csv"}:
        text = read_table(path)
    elif suffix in {".html", ".htm"}:
        text = read_html(path)
    elif suffix == ".docx":
        text = read_docx(path)
    elif suffix == ".doc":
        text = read_doc(path)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")
    return normalize_text(text)[:MAX_DOC_CHARS]


def normalize_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def bucket_documents(documents: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    buckets: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for doc in documents:
        buckets[doc.get("category", "unknown")].append(doc)
        buckets[f"format:{doc.get('format', 'unknown')}"].append(doc)
    return buckets


def choose_sources(
    documents: list[dict[str, Any]],
    buckets: dict[str, list[dict[str, Any]]],
    category: str,
    index: int,
    rng: random.Random,
) -> list[dict[str, Any]]:
    if category == "entity_relation":
        pool = buckets["prd"] + buckets["functional_req"] + buckets["tech_blog"]
        return [pool[index % len(pool)]]
    if category == "multi_hop":
        first_pool = buckets["prd"] + buckets["functional_req"] or documents
        second_pool = buckets["tech_blog"] + buckets["non_functional_req"] or documents
        first = first_pool[index % len(first_pool)]
        second = second_pool[(index * 3) % len(second_pool)]
        if first["doc_id"] == second["doc_id"]:
            second = documents[(index + 1) % len(documents)]
        return [first, second]
    if category == "summary":
        pool = buckets["pure_srs"] + buckets["github_prd"] + buckets["tech_blog"]
        return [pool[index % len(pool)]]
    if category == "distractor":
        return []

    pool = buckets["functional_req"] + buckets["non_functional_req"] + buckets["prd"] + buckets["tech_blog"]
    if not pool:
        pool = documents
    return [pool[index % len(pool)]]


def format_doc_context(doc: dict[str, Any], text: str) -> str:
    meta = {
        "doc_id": doc.get("doc_id"),
        "title": doc.get("title"),
        "subset": doc.get("subset"),
        "category": doc.get("category"),
        "language": doc.get("language"),
        "source_url": doc.get("source_url"),
        "source_site": doc.get("source_site"),
    }
    return f"METADATA:\n{json.dumps(meta, ensure_ascii=False)}\n\nCONTENT:\n{text}"


def build_prompt(
    category: str,
    case_id: str,
    sources: list[dict[str, Any]],
    contexts: list[str],
) -> str:
    source_ids = [doc["doc_id"] for doc in sources]
    shared_rules = f"""
你是企业知识库 RAG 离线评测集的出题专家。请基于给定公开文档内容生成 1 条中文评测题。

必须只输出一个 JSON object, 不要 Markdown, 不要解释。
字段必须包含:
- case_id: "{case_id}"
- category: "{category}"
- question: 中文问题
- expected_keywords: 2 到 4 个判别性强的关键词, 禁止使用 "公司"、"文档"、"信息"、"内容" 等空词
- source_doc_ids: {json.dumps(source_ids, ensure_ascii=False)}
- difficulty: "easy" / "medium" / "hard"
- mode: "mix"

题型要求:
- fact_qa: 单文档单点事实查询。
- entity_relation: 必须询问两个实体之间的关系, 例如需求、模块、接口、指标、约束或技术方案之间的关系。
- multi_hop: 必须需要至少 2 份文档共同支撑答案, source_doc_ids 至少 2 个。
- summary: 文档主旨、段落摘要或主题归纳。
- distractor: 答案必须不在 corpus 中, case 里必须额外包含 expected_behavior: "refuse"。

请保证问题自然、像企业内网用户真实会问的问题。expected_keywords 不要照抄大段句子。
"""

    if category == "distractor":
        return shared_rules + """
本题是 distractor。请生成一个与企业项目文档、需求说明、PRD、技术方案或工程知识库语境相关，
但无法从给定 corpus 判断真假的反事实或未收录信息问题。source_doc_ids 必须是空数组。
输出 JSON 必须包含 expected_behavior: "refuse"。
"""

    joined_context = "\n\n--- DOCUMENT ---\n\n".join(contexts)
    return shared_rules + f"\n给定文档如下:\n\n{joined_context}\n"


def extract_json_object(text: str) -> dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.S)
        if not match:
            raise
        return json.loads(match.group(0))


def call_gemini(api_key: str, model: str, prompt: str, timeout_s: int) -> tuple[dict[str, Any], GenerationStats]:
    endpoint = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.2,
            "responseMimeType": "application/json",
        },
    }
    with httpx.Client(timeout=timeout_s) as client:
        response = client.post(endpoint, headers={"x-goog-api-key": api_key}, json=payload)
    if response.status_code >= 400:
        detail = response.text[:800].replace(api_key, "***")
        raise RuntimeError(f"Gemini API error {response.status_code}: {detail}")
    body = response.json()
    usage = body.get("usageMetadata", {})
    stats = GenerationStats(
        prompt_tokens=int(usage.get("promptTokenCount") or 0),
        output_tokens=int(usage.get("candidatesTokenCount") or 0),
        total_tokens=int(usage.get("totalTokenCount") or 0),
    )
    parts = body["candidates"][0]["content"].get("parts", [])
    text = "\n".join(part.get("text", "") for part in parts)
    return extract_json_object(text), stats


def validate_case(case: dict[str, Any], category: str, case_id: str) -> dict[str, Any]:
    case["case_id"] = case_id
    case["category"] = category
    case.setdefault("mode", "mix")

    required = ["case_id", "category", "question", "expected_keywords", "source_doc_ids", "difficulty", "mode"]
    missing = [field for field in required if field not in case]
    if missing:
        raise ValueError(f"missing fields: {missing}")
    if category == "distractor" and case.get("expected_behavior") != "refuse":
        raise ValueError("distractor must include expected_behavior=refuse")
    if category == "multi_hop" and len(case.get("source_doc_ids", [])) < 2:
        raise ValueError("multi_hop must include at least 2 source_doc_ids")
    keywords = case.get("expected_keywords")
    if not isinstance(keywords, list) or not 2 <= len(keywords) <= 4:
        raise ValueError("expected_keywords must contain 2-4 items")
    weak = {"公司", "文档", "信息", "内容", "数据", "情况"}
    if any(str(keyword).strip() in weak for keyword in keywords):
        raise ValueError(f"weak expected_keywords: {keywords}")
    if case.get("difficulty") not in {"easy", "medium", "hard"}:
        case["difficulty"] = "medium"
    return case


def generate_cases(args: argparse.Namespace) -> tuple[dict[str, Any], Counter[str], list[str], GenerationStats]:
    api_key = load_env_key(REPO_ROOT / ".env")
    manifest_path = args.corpus_dir / "manifest.json"
    documents = load_manifest(manifest_path)
    buckets = bucket_documents(documents)
    rng = random.Random(args.seed)
    counters: Counter[str] = Counter()
    failures: list[str] = []
    usage = GenerationStats()
    queries: list[dict[str, Any]] = []

    for category, quota in SAMPLE_QUOTAS.items():
        for i in range(quota):
            case_id = f"{CASE_PREFIX[category]}-{i + 1:03d}"
            sources = choose_sources(documents, buckets, category, i, rng)
            contexts = []
            try:
                for doc in sources:
                    contexts.append(format_doc_context(doc, read_document(args.raw_root, doc)))
            except Exception as exc:
                failures.append(f"{case_id}: read failed: {exc}")
                continue

            prompt = build_prompt(category, case_id, sources, contexts)
            last_error = ""
            for attempt in range(1, args.retries + 1):
                try:
                    raw_case, stats = call_gemini(api_key, args.model, prompt, args.timeout_s)
                    usage.prompt_tokens += stats.prompt_tokens
                    usage.output_tokens += stats.output_tokens
                    usage.total_tokens += stats.total_tokens
                    case = validate_case(raw_case, category, case_id)
                    if category != "distractor":
                        expected_ids = [doc["doc_id"] for doc in sources]
                        case["source_doc_ids"] = expected_ids
                    else:
                        case["source_doc_ids"] = []
                    queries.append(case)
                    counters[category] += 1
                    print(f"[ok] {case_id} {category}: {case['question']}", flush=True)
                    break
                except Exception as exc:
                    last_error = str(exc)
                    print(f"[retry {attempt}/{args.retries}] {case_id}: {last_error}", flush=True)
                    time.sleep(min(2 * attempt, 6))
            else:
                failures.append(f"{case_id}: {last_error}")

    result = {
        "version": VERSION,
        "generated_by": args.model,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "corpus_manifest": str(manifest_path.relative_to(REPO_ROOT)),
        "queries": queries,
    }
    return result, counters, failures, usage


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--corpus-dir", type=Path, default=DEFAULT_CORPUS_DIR)
    parser.add_argument("--raw-root", type=Path, default=DEFAULT_RAW_ROOT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--model", default=DEFAULT_MODEL)
    parser.add_argument("--seed", type=int, default=20260519)
    parser.add_argument("--retries", type=int, default=3)
    parser.add_argument("--timeout-s", type=int, default=120)
    return parser.parse_args()


def main() -> int:
    start = time.perf_counter()
    args = parse_args()
    args.corpus_dir = args.corpus_dir.resolve()
    args.raw_root = args.raw_root.resolve()
    args.output = args.output.resolve()

    try:
        result, counters, failures, usage = generate_cases(args)
    except Exception as exc:
        print(f"[fatal] {exc}", file=sys.stderr)
        return 1

    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(result, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    elapsed = time.perf_counter() - start
    print("\n=== generation summary ===")
    print(f"output: {args.output}")
    print(f"total_cases: {len(result['queries'])}")
    print(f"by_category: {dict(counters)}")
    print(f"failures: {len(failures)}")
    for failure in failures:
        print(f"- {failure}")
    print(
        "token_usage: "
        f"input={usage.prompt_tokens}, output={usage.output_tokens}, total={usage.total_tokens}"
    )
    print(f"elapsed_s: {elapsed:.2f}")
    return 0 if len(result["queries"]) == sum(SAMPLE_QUOTAS.values()) else 2


if __name__ == "__main__":
    raise SystemExit(main())
